import unittest
import tempfile
from pathlib import Path
from unittest import mock

from lxml import etree

from diff_engine import diff_trees
from spec_parser import (
    _build_tree,
    _clause_level,
    _convert_emf_to_png,
    _element_text,
    _extract_clause_id,
    _heading_level,
    _is_emf_content,
    _merge_images,
    _table_to_text,
    _vector_preview_path,
    convert_doc_to_docx,
)


class ClauseIdParsingTests(unittest.TestCase):
    def test_extracts_annex_clause_ids(self):
        self.assertEqual(
            _extract_clause_id(
                "D.5 Support for keeping UE in CM-CONNECTED state "
                "in overlay network when accessing services via NWu"
            ),
            "D.5",
        )
        self.assertEqual(_extract_clause_id("D.7.1 Network initiated QoS"), "D.7.1")
        self.assertEqual(
            _extract_clause_id("Annex D (informative): deployment options"),
            "Annex D",
        )

    def test_normalizes_spaces_inside_numeric_clause_ids(self):
        self.assertEqual(_extract_clause_id("5.2. 1 General"), "5.2.1")
        self.assertEqual(_extract_clause_id("4 . 3 Security aspects"), "4.3")
        self.assertEqual(_extract_clause_id("5.35A.1 General"), "5.35A.1")
        self.assertIsNone(_extract_clause_id("3GPP EPS architecture"))
        self.assertIsNone(_extract_clause_id("5GS features"))
        self.assertEqual(
            _extract_clause_id(
                "There are no custom operations in this release.5.27.4\tNotifications"
            ),
            "5.27.4",
        )
        self.assertIsNone(_extract_clause_id("Architecture version 23.501"))

    def test_title_spacing_change_is_a_modification_not_add_delete(self):
        old_tree = _build_tree([
            (
                1,
                "D.5 Support for keeping UE in CM-CONNECTED state "
                "in overlay network when accessing services via NWu",
                ["old wording"],
            )
        ])
        new_tree = _build_tree([
            (
                1,
                "D.5 Support for keeping UE in CM-CONNECTED state "
                "in overlay network when accessing services via  NWu",
                ["new wording"],
            )
        ])

        result = diff_trees(old_tree, new_tree)

        self.assertEqual(len(result), 1)
        self.assertEqual(result[0]["id"], "D.5")
        self.assertEqual(result[0]["status"], "modified")
        self.assertEqual(result[0]["old_body"], "old wording")
        self.assertEqual(result[0]["new_body"], "new wording")

    def test_number_prefix_overrides_an_out_of_order_heading(self):
        tree = _build_tree([
            (1, "5 Service APIs", []),
            (2, "5.6 Analytics API", []),
            (3, "5.6.1 Resources", []),
            (3, "5.6.0 Introduction", ["new introductory text"]),
            (4, "5.6.1.1 Overview", ["resource overview"]),
            (4, "5.6.1.2 Subscriptions", ["resource definition"]),
        ])

        api = tree[0]["children"][0]
        resources, introduction = api["children"]
        self.assertEqual(resources["id"], "5.6.1")
        self.assertEqual(
            [node["id"] for node in resources["children"]],
            ["5.6.1.1", "5.6.1.2"],
        )
        self.assertEqual(introduction["id"], "5.6.0")
        self.assertEqual(introduction["children"], [])

    def test_distant_number_prefix_does_not_override_document_context(self):
        tree = _build_tree([
            (1, "5 Service APIs", []),
            (2, "5.4 Earlier API", []),
            (3, "5.4.2 Resources", []),
            (4, "5.4.2.2 Subscription", []),
            (5, "5.4.2.2.3 Operations", []),
            (2, "5.20 Current API", []),
            (3, "5.20.4 Resources", []),
            (4, "5.20.4.2 Subscription", []),
            (5, "5.20.4.2.3 Operations", []),
            # Real specifications occasionally retain a copied clause number.
            (6, "5.4.2.2.3.2 Notification via Websocket", ["current API body"]),
        ])

        current_operations = tree[0]["children"][1]["children"][0]["children"][0]["children"][0]
        self.assertEqual(
            [node["id"] for node in current_operations["children"]],
            ["5.4.2.2.3.2"],
        )

    def test_text_runs_do_not_create_false_spaces(self):
        paragraph = etree.fromstring(
            b'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            b'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
            b'<w:r><w:t>5</w:t></w:r><w:r><w:t>GS</w:t></w:r>'
            b'<w:r><w:tab/></w:r><w:r><w:t>value</w:t><w:br/><w:t>non</w:t>'
            b'<w:noBreakHyphen/><w:t>breaking</w:t></w:r><m:oMath><m:r><m:t>=E</m:t></m:r></m:oMath>'
            b'</w:p>'
        )

        self.assertEqual(_element_text(paragraph), "5GS\tvalue\nnon-breaking=E")

    def test_tables_keep_rows_rectangular_and_flatten_cell_paragraphs(self):
        table = etree.fromstring(
            b'<w:tbl xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            b'<w:tr><w:tc><w:p><w:r><w:t>Information element</w:t></w:r></w:p></w:tc>'
            b'<w:tc><w:p><w:r><w:t>Description</w:t></w:r></w:p></w:tc></w:tr>'
            b'<w:tr><w:tc><w:p><w:r><w:t>First</w:t></w:r></w:p>'
            b'<w:p><w:r><w:t>second</w:t></w:r></w:p></w:tc>'
            b'<w:tc><w:p><w:r><w:t>A | B</w:t></w:r></w:p></w:tc></w:tr>'
            b'</w:tbl>'
        )
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

        self.assertEqual(
            _table_to_text(table, ns),
            "| Information element | Description |\n"
            "| First\\nsecond | A \\| B |",
        )

    def test_tables_expand_simple_horizontal_spans(self):
        table = etree.fromstring(
            b'<w:tbl xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            b'<w:tr><w:tc><w:tcPr><w:gridSpan w:val="2"/></w:tcPr>'
            b'<w:p><w:r><w:t>Wide heading</w:t></w:r></w:p></w:tc></w:tr>'
            b'<w:tr><w:tc><w:p><w:r><w:t>Left</w:t></w:r></w:p></w:tc>'
            b'<w:tc><w:p><w:r><w:t>Right</w:t></w:r></w:p></w:tc></w:tr>'
            b'</w:tbl>'
        )
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

        self.assertEqual(
            _table_to_text(table, ns),
            "| Wide heading |  |\n| Left | Right |",
        )

    def test_images_are_merged_by_unique_source_position(self):
        tree = _build_tree([
            (1, "5.1 First duplicate", [], 10),
            (1, "5.1 Second duplicate", [], 20),
        ])

        _merge_images(tree, {20: [{"id": "figure2", "sha256": "abc"}]})

        self.assertNotIn("images", tree[0])
        self.assertEqual(tree[1]["images"][0]["id"], "figure2")
        self.assertNotIn("_source_key", tree[0])
        self.assertNotIn("_source_key", tree[1])

    def test_heading_level_accepts_docx_style_ids(self):
        self.assertEqual(_heading_level("Heading3"), 3)
        self.assertEqual(_heading_level("Heading 4"), 4)
        self.assertEqual(_heading_level("heading8"), 8)
        self.assertEqual(_clause_level("Annex D", "Heading8"), 1)
        self.assertEqual(_clause_level("D.1", "Heading1"), 2)
        self.assertEqual(_clause_level("D.7.2", "Heading2"), 3)

    def test_parses_custom_styles_derived_from_a_heading(self):
        from docx import Document
        from docx.enum.style import WD_STYLE_TYPE
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from spec_parser import parse_spec

        with tempfile.TemporaryDirectory() as tempdir:
            path = Path(tempdir) / "derived-heading.docx"
            document = Document()
            title_style = document.styles.add_style("TT", WD_STYLE_TYPE.PARAGRAPH)
            title_style.base_style = document.styles["Heading 1"]
            contents = document.add_paragraph("Contents")
            contents.style = title_style
            document.add_heading("5 Parent", level=1)
            derived = document.styles.add_style("H6", WD_STYLE_TYPE.PARAGRAPH)
            derived.base_style = document.styles["Heading 5"]
            paragraph = document.add_paragraph("5.1 Derived heading")
            paragraph.style = derived
            document.add_paragraph("Child body")

            outline_heading = document.add_paragraph("5.2 Explicit outline heading")
            outline = OxmlElement("w:outlineLvl")
            outline.set(qn("w:val"), "2")
            outline_heading._p.get_or_add_pPr().append(outline)
            document.add_paragraph("Outlined child body")

            body_outline = document.add_paragraph("5.3 Outline level nine is body")
            body_level = OxmlElement("w:outlineLvl")
            body_level.set(qn("w:val"), "9")
            body_outline._p.get_or_add_pPr().append(body_level)
            document.save(path)

            parsed = parse_spec(path, spec_number="23.501", version="20.1.0")

        self.assertEqual([node["id"] for node in parsed["clauses"]], ["5"])
        child = parsed["clauses"][0]["children"][0]
        self.assertEqual(child["id"], "5.1")
        self.assertEqual(child["body"], "Child body")
        outlined = parsed["clauses"][0]["children"][1]
        self.assertEqual(outlined["id"], "5.2")
        self.assertEqual(
            outlined["body"],
            "Outlined child body\n5.3 Outline level nine is body",
        )

    def test_combines_ordered_multi_part_documents(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as tempdir:
            root = Path(tempdir)
            main_path = root / "spec_1_main.docx"
            annex_path = root / "spec_2_annex.docx"

            main = Document()
            main.core_properties.subject = "Combined specification"
            main.add_heading("1 Scope", level=1)
            main.add_paragraph("Main body")
            main.add_heading("1.1 Detail", level=2)
            main.save(main_path)

            annex = Document()
            annex.add_heading("Annex A (informative): Examples", level=1)
            annex.add_heading("A.1 General", level=2)
            annex.add_paragraph("Annex body")
            annex.save(annex_path)

            from spec_parser import parse_spec
            parsed = parse_spec(
                [main_path, annex_path], spec_number="29.522", version="19.0.0"
            )

        self.assertEqual(parsed["title"], "Combined specification")
        self.assertEqual([node["id"] for node in parsed["clauses"]], ["1", "Annex A"])
        self.assertEqual(parsed["clauses"][0]["children"][0]["id"], "1.1")
        self.assertEqual(parsed["clauses"][1]["children"][0]["id"], "A.1")

    def test_caller_version_overrides_stale_document_metadata(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as tempdir:
            path = Path(tempdir) / "stale.docx"
            document = Document()
            document.core_properties.subject = "Example service (Release 18)"
            document.add_paragraph("3GPP TS 29.522 V18.6.0")
            document.add_heading("1 Scope", level=1)
            document.save(path)

            from spec_parser import parse_spec
            parsed = parse_spec(path, spec_number="29.522", version="19.0.0")

        self.assertEqual(parsed["spec_number"], "29.522")
        self.assertEqual(parsed["version"], "19.0.0")
        self.assertEqual(parsed["release"], 19)
        self.assertEqual(parsed["title"], "Example service (Release 19)")

    def test_converts_vector_figures_with_libreoffice_and_keeps_originals(self):
        with tempfile.TemporaryDirectory() as tempdir:
            root = Path(tempdir)
            images = [{"src": f"figure{index}.emf"} for index in range(8)]
            for image in images:
                (root / image["src"]).write_bytes(b"vector")

            def fake_libreoffice(paths):
                for path in paths:
                    _vector_preview_path(path).write_bytes(b"png")

            with mock.patch(
                "spec_parser._convert_vectors_with_libreoffice",
                side_effect=fake_libreoffice,
            ) as libreoffice:
                _convert_emf_to_png(root, {"1": images})

            libreoffice.assert_called_once()
            self.assertTrue(
                all(image["src"].endswith(".preview.png") for image in images)
            )
            self.assertTrue(
                all(image["original_src"].endswith(".emf") for image in images)
            )
            self.assertEqual(len(list(root.glob("*.emf"))), len(images))

    def test_inkscape_fallback_corrects_mislabeled_emf_extension(self):
        with tempfile.TemporaryDirectory() as tempdir:
            root = Path(tempdir)
            source = root / "figure.wmf"
            header = bytearray(44)
            header[:4] = b"\x01\x00\x00\x00"
            header[40:44] = b" EMF"
            source.write_bytes(header)
            images = [{"src": source.name}]
            commands = []

            def fake_inkscape(command, **_kwargs):
                commands.append(command)
                for value in command:
                    path = Path(value)
                    if path.name.endswith(".preview.emf"):
                        path.with_suffix(".png").write_bytes(b"png")
                return mock.Mock(returncode=0, stderr="", stdout="")

            with mock.patch(
                "spec_parser._convert_vectors_with_libreoffice"
            ), mock.patch("spec_parser.subprocess.run", side_effect=fake_inkscape):
                _convert_emf_to_png(root, {"1": images})

            self.assertTrue(_is_emf_content(source))
            self.assertTrue(source.exists())
            self.assertEqual(images[0]["src"], "figure.preview.png")
            self.assertTrue(
                any(
                    str(value).endswith("figure.preview.emf")
                    for value in commands[0]
                )
            )

    def test_libreoffice_conversion_uses_an_isolated_profile(self):
        with tempfile.TemporaryDirectory() as tempdir:
            source = Path(tempdir) / "legacy.doc"
            source.write_bytes(b"legacy")
            commands = []

            def fake_libreoffice(command, **_kwargs):
                commands.append(command)
                outdir = Path(command[command.index("--outdir") + 1])
                (outdir / "legacy.docx").write_bytes(b"converted")
                return mock.Mock(returncode=0, stderr="")

            with mock.patch(
                "spec_parser.subprocess.run", side_effect=fake_libreoffice
            ):
                converted = convert_doc_to_docx(source)

            self.assertEqual(converted, source.with_suffix(".docx"))
            self.assertEqual(converted.read_bytes(), b"converted")
            profile_args = [
                arg for arg in commands[0]
                if arg.startswith("-env:UserInstallation=file:")
            ]
            self.assertEqual(len(profile_args), 1)

    def test_libreoffice_conversion_retries_a_transient_failure(self):
        with tempfile.TemporaryDirectory() as tempdir:
            source = Path(tempdir) / "legacy.doc"
            source.write_bytes(b"legacy")
            attempts = 0

            def flaky_libreoffice(command, **_kwargs):
                nonlocal attempts
                attempts += 1
                if attempts == 1:
                    return mock.Mock(returncode=1, stderr="temporary startup failure")
                outdir = Path(command[command.index("--outdir") + 1])
                (outdir / "legacy.docx").write_bytes(b"converted")
                return mock.Mock(returncode=0, stderr="")

            with mock.patch(
                "spec_parser.subprocess.run", side_effect=flaky_libreoffice
            ):
                converted = convert_doc_to_docx(source)

            self.assertEqual(converted.read_bytes(), b"converted")

        self.assertEqual(attempts, 2)


if __name__ == "__main__":
    unittest.main()
