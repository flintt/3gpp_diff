"""
Parse 3GPP Word (.doc/.docx) specification documents into structured clause trees.
"""
import re
import subprocess
import tempfile
from pathlib import Path


def parse_spec(doc_path: Path, spec_number: str = None, version: str = None) -> dict:
    """
    Parse a 3GPP spec .doc or .docx into a structured document.
    Accepts optional spec_number and version from the caller as a fallback
    (needed when document metadata is lost, e.g. after LibreOffice conversion).

    Returns:
    {
        "title": "System architecture for the 5G System (5GS)",
        "spec_number": "23.501",
        "version": "18.0.0",
        "release": 18,
        "clauses": [
            {
                "id": "4",
                "title": "Architecture model and concepts",
                "level": 1,
                "body": "text...",
                "children": [...]
            }
        ]
    }
    """
    ext = doc_path.suffix.lower()

    if ext == ".docx":
        return _parse_docx(doc_path, spec_number, version)
    elif ext == ".doc":
        return _parse_doc_via_libreoffice(doc_path, spec_number, version)
    else:
        raise ValueError(f"Unsupported format: {ext}")


def convert_doc_to_docx(doc_path: Path) -> Path:
    """Convert .doc to .docx using LibreOffice headless."""
    output_dir = doc_path.parent
    # LibreOffice can't overwrite the same file, so use a temp dir
    with tempfile.TemporaryDirectory() as tmpdir:
        result = subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                tmpdir,
                str(doc_path),
            ],
            capture_output=True,
            text=True,
            timeout=120,
        )
        if result.returncode != 0:
            raise RuntimeError(
                f"LibreOffice conversion failed: {result.stderr[:500]}"
            )

        # Find the output file
        tmpfiles = list(Path(tmpdir).glob("*.docx"))
        if not tmpfiles:
            raise FileNotFoundError("LibreOffice produced no output")

        # Move to cache dir
        out_path = doc_path.with_suffix(".docx")
        import shutil

        shutil.move(str(tmpfiles[0]), str(out_path))
        return out_path


def _parse_doc_via_libreoffice(doc_path: Path, spec_number: str = None, version: str = None) -> dict:
    """Parse old .doc by converting to .docx first, then parsing."""
    docx_path = convert_doc_to_docx(doc_path)
    return _parse_docx(docx_path, spec_number, version)


def _parse_docx(docx_path: Path, spec_number: str = None, version: str = None) -> dict:
    """Parse a .docx file into structured document with image references."""
    from docx import Document

    doc = Document(str(docx_path))

    # Pre-compute paragraph data once to share between metadata extraction
    # and clause tree building (avoids redundant slow python-docx access)
    para_data = [(para.style.name, para.text.strip()) for para in doc.paragraphs]

    metadata = _extract_metadata(doc, spec_number, version, _para_data=para_data)

    clauses, para_clause_map = _build_clause_tree(doc, _para_data=para_data)

    # Extract images and merge into clause tree
    if para_clause_map:
        cache_dir = Path("cache") / "images" / (metadata.get("spec_number", "_")) / (metadata.get("version", "_"))
        clause_images = _extract_images_from_docx(docx_path, cache_dir, para_clause_map)
        if clause_images:
            _merge_images(clauses, clause_images)

    return {
        "title": metadata.get("title", ""),
        "spec_number": metadata.get("spec_number", ""),
        "version": metadata.get("version", ""),
        "release": metadata.get("release", 0),
        "clauses": clauses,
    }


def _extract_metadata(doc, spec_number: str = None, version: str = None, _para_data: list = None) -> dict:
    """Extract spec number, version, title from the document."""
    meta = {
        "title": "",
        "spec_number": spec_number or "",
        "version": version or "",
        "release": 0,
    }
    if version:
        major = version.split(".")[0]
        try:
            meta["release"] = int(major)
        except ValueError:
            pass

    # Try core_properties.subject for the canonical title (reliable for native .docx)
    try:
        props = doc.core_properties
        if props.subject:
            meta["title"] = props.subject
    except Exception:
        pass

    # Search for 3GPP version pattern in paragraphs (early exit on match)
    paras = _para_data if _para_data is not None else [(None, p.text.strip()) for p in doc.paragraphs]
    for _style, t in paras:
        if not t:
            continue
        m = re.search(r"3GPP\s+TS\s+(\d+\.\d+)\s+V(\d+)\.(\d+)\.(\d+)", t)
        if m:
            meta["spec_number"] = m.group(1)
            meta["version"] = f"{m.group(2)}.{m.group(3)}.{m.group(4)}"
            meta["release"] = int(m.group(2))
            break

    if not meta["title"]:
        if meta["spec_number"] and meta["version"]:
            meta["title"] = f"3GPP TS {meta['spec_number']} v{meta['version']} (Release {meta['release']})"

    return meta


def _build_clause_tree(doc, _para_data: list = None) -> tuple:
    """Build hierarchical clause tree from Word heading styles.
    
    Returns:
        tuple: (clauses, para_clause_map)
            clauses: list of clause tree nodes
            para_clause_map: dict mapping paragraph index -> clause_id
    """
    # Pre-compute paragraph data once if not already provided
    if _para_data is None:
        para_data = [(para.style.name, para.text.strip()) for para in doc.paragraphs]
    else:
        para_data = _para_data
    n = len(para_data)

    para_clause_map = {}
    entries = []

    para_idx = 0
    while para_idx < n:
        style, text = para_data[para_idx]

        if "heading" in style.lower() and text:
            level = _heading_level(style)
            tmp = _extract_clause_id(text)
            clause_id = tmp if tmp else (text.split("\t")[0].strip() if "\t" in text else text)

            para_clause_map[para_idx] = clause_id

            body_texts = []
            para_idx += 1
            while para_idx < n:
                next_style, next_text = para_data[para_idx]
                if "heading" in next_style.lower() and next_text:
                    break
                if not next_style.startswith("toc"):
                    para_clause_map[para_idx] = clause_id
                if next_text and not next_style.startswith("toc"):
                    body_texts.append(next_text)
                para_idx += 1

            entries.append((level, text, body_texts))
        else:
            para_idx += 1

    root = _build_tree(entries)
    return root, para_clause_map


def _heading_level(style_name: str) -> int:
    """Extract heading level from Word style name.
    'Heading 1' -> 1, 'Heading 2' -> 2, 'Heading 3' -> 3
    Default: 1
    """
    if style_name.startswith("Heading "):
        try:
            return int(style_name.split()[1])
        except (ValueError, IndexError):
            return 1
    return 1


def _extract_images_from_docx(docx_path, cache_dir, para_clause_map):
    """Extract images from .docx and associate with clauses by paragraph index.
    
    Returns dict mapping clause_id -> list of image info dicts.
    """
    import zipfile
    from lxml import etree

    NS = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        'v': 'urn:schemas-microsoft-com:vml',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    }
    R_ATTR = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}'

    clause_images = {}

    try:
        with zipfile.ZipFile(str(docx_path), 'r') as z:
            # Parse relationships to build rel_id -> target path mapping for images
            rels_xml = etree.parse(z.open('word/_rels/document.xml.rels'))
            rel_map = {}
            for rel in rels_xml.xpath('//*[local-name()="Relationship"]'):
                rel_id = rel.get('Id')
                target = rel.get('Target')
                rel_type = rel.get('Type', '')
                if 'relationships/image' in rel_type:
                    rel_map[rel_id] = target

            if not rel_map:
                return clause_images

            # Parse document body paragraphs
            doc_xml = etree.parse(z.open('word/document.xml'))
            body = doc_xml.xpath('/w:document/w:body', namespaces=NS)
            if not body:
                return clause_images
            body = body[0]
            paragraphs = body.xpath('w:p', namespaces=NS)

            for para_idx, para in enumerate(paragraphs):
                if para_idx not in para_clause_map:
                    continue

                clause_id = para_clause_map[para_idx]
                seen_targets = set()

                # VML images (<v:imagedata r:id="..."> inside <w:object> or <w:pict>)
                for imagedata in para.xpath('.//v:imagedata', namespaces=NS):
                    # Some v:imagedata elements reference OLE objects, not images.
                    # Only process those with valid r:id pointing to an image relationship.
                    rel_id = imagedata.get(f'{R_ATTR}id')
                    if rel_id and rel_id in rel_map:
                        target = rel_map[rel_id]
                        if target not in seen_targets:
                            seen_targets.add(target)
                            info = _save_image(z, target, cache_dir, clause_id)
                            if info:
                                clause_images.setdefault(clause_id, []).append(info)

                # DrawingML images (<a:blip r:embed="...">)
                for blip in para.xpath('.//a:blip', namespaces=NS):
                    embed = blip.get(f'{R_ATTR}embed')
                    if embed and embed in rel_map:
                        target = rel_map[embed]
                        if target not in seen_targets:
                            seen_targets.add(target)
                            info = _save_image(z, target, cache_dir, clause_id)
                            if info:
                                clause_images.setdefault(clause_id, []).append(info)

    except Exception as e:
        import warnings
        warnings.warn(f"Image extraction failed for {docx_path}: {e}")

    # Batch-convert EMF/WMF to PNG for browser compatibility
    if clause_images:
        _convert_emf_to_png(Path(cache_dir), clause_images)

    return clause_images


def _convert_emf_to_png(cache_dir, clause_images):
    """Convert all EMF/WMF images in cache_dir to PNG using Inkscape."""
    emf_files = list(cache_dir.glob("*.emf")) + list(cache_dir.glob("*.wmf"))
    if not emf_files:
        return

    import subprocess
    emf_count = len(emf_files)
    for i, emf_path in enumerate(emf_files):
        try:
            png_path = emf_path.with_suffix(".png")
            if png_path.exists():
                continue
            subprocess.run(
                ["inkscape", str(emf_path), "--export-filename", str(png_path)],
                capture_output=True,
                timeout=60,
            )
        except Exception:
            continue

    for emf_path in emf_files:
        png_path = emf_path.with_suffix(".png")
        if png_path.exists():
            emf_path.unlink()

    for images in clause_images.values():
        for img in images:
            src = img["src"]
            if src.lower().endswith((".emf", ".wmf")):
                png_name = Path(src).stem + ".png"
                if (cache_dir / png_name).exists():
                    img["src"] = png_name


def _save_image(zip_file, target, cache_dir, clause_id):
    """Extract image from ZIP to cache and return image info dict."""
    try:
        img_path = f'word/{target}'
        img_data = zip_file.read(img_path)
        img_filename = target.split('/')[-1]

        cache_dir = Path(cache_dir)
        cache_dir.mkdir(parents=True, exist_ok=True)
        cache_path = cache_dir / img_filename
        if not cache_path.exists():
            cache_path.write_bytes(img_data)

        src = img_filename
        # If this image was previously converted from EMF/WMF to PNG
        # (file in cache is PNG but original ZIP name is EMF), use PNG ref
        if img_filename.lower().endswith(('.emf', '.wmf')):
            png_path = cache_dir / (Path(img_filename).stem + '.png')
            if png_path.exists():
                src = png_path.name

        return {
            'id': img_filename,
            'src': src,
            'alt': f'Figure in clause {clause_id}',
        }
    except Exception:
        return None


def _merge_images(clauses, clause_images):
    """Recursively merge image references into clause tree nodes."""
    for node in clauses:
        cid = node.get("id", "")
        if cid in clause_images:
            node["images"] = clause_images[cid]
        _merge_images(node.get("children", []), clause_images)


def _build_tree(entries: list, base_level: int = 1) -> list:
    """Convert flat list of (level, heading, body) into nested tree.
    Only includes entries at level >= base_level.
    """
    tree = []
    stack = []  # stack of parent nodes

    for level, heading, body in entries:
        # Skip entries that don't have clause-style IDs (Foreword, etc.)
        clause_id = _extract_clause_id(heading)

        node = {
            "id": clause_id or heading.split("\t")[0].strip(),
            "title": heading.split("\t")[-1].strip() if "\t" in heading else heading,
            "raw_heading": heading,
            "level": level,
            "body": "\n".join(body),
            "changed": False,
            "children": [],
        }

        # Pop stack until we find the right parent
        while stack and stack[-1]["level"] >= level:
            stack.pop()

        if stack:
            stack[-1]["children"].append(node)
        else:
            tree.append(node)

        stack.append(node)

    return tree


def _extract_clause_id(heading: str) -> str:
    """Extract clause number from heading text.
    '4.2.3\tNon-roaming reference architecture' -> '4.2.3'
    '4.2.3 Non-roaming reference architecture' -> '4.2.3'
    Returns None if no clause number found.
    """
    # Tab-separated
    if "\t" in heading:
        candidate = heading.split("\t")[0].strip()
        if re.match(r"^\d+(\.\d+)*$", candidate):
            return candidate

    # First word might be clause number
    first_word = heading.split()[0] if heading.split() else ""
    if re.match(r"^\d+(\.\d+)*$", first_word):
        return first_word

    return None


def clause_count(clauses: list) -> int:
    """Count total clauses (recursive)."""
    count = 0
    for c in clauses:
        count += 1
        if c.get("children"):
            count += clause_count(c["children"])
    return count


if __name__ == "__main__":
    import json
    import sys

    path = Path(sys.argv[1])
    doc = parse_spec(path)
    print(f"Title: {doc['title']}")
    print(f"Spec: {doc['spec_number']} v{doc['version']}")
    print(f"Total clauses: {clause_count(doc['clauses'])}")

    # Print top-level clauses
    for c in doc["clauses"]:
        n = clause_count([c])
        print(f"  [{c['id']}] {c['title']} ({n} clauses, {len(c.get('body',''))} chars)")
